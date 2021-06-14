# mistune docx renderer code

# This code is from the mistune python package, which provides a command line utility
# to convert Markdown to .docx. Unfortunately, the public implementation cannot be
# imported as a module. 

# https://github.com/mjanv/mistune-docx/blob/24e8e64fe096f07079c785ab9a359eb342d5ed2b/generate_doc.py

import os
import re
import itertools

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import mistune

class MathBlockGrammar(mistune.BlockGrammar):
    block_math = re.compile(r"^\$\$(.*?)\$\$", re.DOTALL)


class MathBlockLexer(mistune.BlockLexer):
    default_rules = ['block_math'] + mistune.BlockLexer.default_rules

    def __init__(self, rules=None, **kwargs):
        if rules is None:
            rules = MathBlockGrammar()
        super(MathBlockLexer, self).__init__(rules, **kwargs)

    def parse_block_math(self, m):
        """Parse a $$math$$ block"""
        self.tokens.append({'type': 'block_math', 'text': m.group(1)})


class MarkdownWithMath(mistune.Markdown):
    def __init__(self, renderer, **kwargs):
        kwargs['block'] = MathBlockLexer
        super(MarkdownWithMath, self).__init__(renderer, **kwargs)

    def output_block_math(self):
        return self.renderer.block_math(self.token['text'])


class PythonDocxRenderer(mistune.Renderer):
    def __init__(self, **kwds):
        super(PythonDocxRenderer, self).__init__(**kwds)
        self.table_memory = []
        self.img_counter = 0

    def header(self, text, level, raw):
        return "p = document.add_heading('', %d)\n" % (level - 1) + text

    def paragraph(self, text):
        if 'add_picture' in text:
            return text
        add_break = '' if text.endswith(':")\n') else 'p.add_run().add_break()'
        return '\n'.join(('p = document.add_paragraph()', text, add_break)) + '\n'

    def list(self, body, ordered):
        return body + '\np.add_run().add_break()\n'

    def list_item(self, text, style=None): #style='BasicUserList'
        if not style:
            return '\n'.join(("p = document.add_paragraph('')", text))
        else:
            return '\n'.join(("p = document.add_paragraph('', style = '%s')" % (style), text))

    def table(self, header, body, style=None): # style='BasicUserTable'
        number_cols = header.count('\n') - 2
        number_rows = int(len(self.table_memory) / number_cols)
        cells = ["table.rows[%d].cells[%d].paragraphs[0]%s\n" 
                 % (i, j, self.table_memory.pop(0)[1:]) for i, j in itertools.product(range(number_rows), range(number_cols))]
        if style is None:
            return '\n'.join(["table = document.add_table(rows=%d, cols=%d)" 
                              % (number_rows, number_cols)] + cells) + 'document.add_paragraph().add_run().add_break()\n'
        else:
            return '\n'.join(["table = document.add_table(rows=%d, cols=%d, style = '%s')" 
                              % (number_rows, number_cols, style)] + cells) + 'document.add_paragraph().add_run().add_break()\n'

    def table_cell(self, content, **flags):
        self.table_memory.append(content)
        return content

    # SPAN LEVEL
    def text(self, text):
        return "p.add_run(\"\"\"%s\"\"\")\n" % text

    def emphasis(self, text):
        return text[:-1] + '.italic = True\n'

    def double_emphasis(self, text):
        return text[:-1] + '.bold = True\n'

    def block_code(self, code, language):
        code = code.replace('\n', '\\n')
        return "p = document.add_paragraph()\np.add_run(\"%s\")\np.style = 'BasicUserQuote'\np.add_run().add_break()\n" % code

    def link(self, link, title, content):
        return "%s (%s)" % (content, link)

    def image(self, src, title, alt_text, width_cm=15):
#  if src.startswith('data:'):
#      from datatools_bdh.data_uri import data_uri_to_bytes
#      run.add_picture(data_uri_to_bytes(src))
        return '\n'.join((
            "p = document.add_paragraph()",
            "p.alignment = WD_ALIGN_PARAGRAPH.CENTER",
            "p.space_after = Pt(18)",
            "run = p.add_run()",
            ("from datatools_bdh.data_uri import data_uri_to_bytes; import io\n"
                "run.add_picture(io.BytesIO(data_uri_to_bytes(\'%s\')), width=Cm(%s))" % (src, width_cm)
                if src.startswith('data:') 
                else "run.add_picture(\'%s\')" % src if "tmp" in src else "run.add_picture(\'%s\', width=Cm(%s))" % (src, width_cm)),
            "run.add_break()",
            "run.add_text(\'%s\')" % alt_text,
            "run.font.italic = True",
            "run.add_break()"
            )) + '\n'

    def hrule(self):
        return "document.add_page_break()\n"

    def block_math(self, text):
        import sympy
        if not os.path.exists('tmp'):
            os.makedirs('tmp')
        filename = 'tmp/tmp%d.png' % self.img_counter
        self.img_counter = self.img_counter + 1
        sympy.preview(r'$$%s$$' % text, output='png', viewer='file', filename=filename, euler=False)
        return self.image(filename, None, "Equation " + str(self.img_counter - 1))

# ----------------------------------------------------------------------------
class MarkdownDocumentBase:
    """Keep markdown source text in member variable T and render to docx before saving."""

    T = None
    renderer = None
    document = None

    def __init__(self):
        self.document = Document()
        self.renderer = PythonDocxRenderer()
        self.T = ""

    def render_markdown(self):
        """Trigger markdown rendering to add components to docx document.
        The internal markdown text variable T is being reset empty."""
        if self.T:
            document = self.document # document variable will be used in exec below
            if isinstance(self.T, str):
                exec(MarkdownWithMath(renderer=self.renderer)(self.T))
            else: # case of T is list of strings, does not apply in current code
                exec(MarkdownWithMath(renderer=self.renderer)('\n'.join(self.T)))
            self.T = ""

    def save_document(self, fname):
        self.render_markdown()
        self.document.save(fname)

    def add_figure(self, fig_uri):
        self.T += f"![]({fig_uri})\n\n"

    def add_text(self, text):
        self.T += text+"\n"

# ----------------------------------------------------------------------------
# app specific custom api, e.g. for different figure drawing backends

class MarkdownDocument(MarkdownDocumentBase):
    """Markdown document with plotly figure support."""

    images_folder = None # location where images are stored
    plotly_fig_opts = None # options to plotly.Figure.write_image call

    def __init__(self, images_folder='images'):
        super().__init__()
        self.images_folder = images_folder
        self.plotly_fig_opts = dict(width=600, height=350, scale=2)

    def add_plotly_figure(self, fig, fig_id=None, caption=None):
        """Write plotly figure to disk under images_folder and insert markdown reference to it.
        Args:
        fig - plotly.graph_objects.Figure object
        fig_id - prefix of .png filename for the figure. If None,
                 the figure will be inserted directly via data: uri,
                 not creating a file in `images_folder`
        caption - Caption to display under the figure
        """
        if not fig_id is None:
            fig_uri = f'{self.images_folder}/{fig_id}.png'
            fig.write_image(fig_uri, **self.plotly_fig_opts)
        else:
            img_bytes = fig.to_image(format="png", **self.plotly_fig_opts)
            from datatools_bdh.data_uri import bytes_to_uri
            fig_uri = bytes_to_uri(img_bytes)
        self.add_figure(fig_uri)
        if not caption is None:
            self.add_text(caption)
